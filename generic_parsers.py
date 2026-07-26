# -*- coding: utf-8 -*-
"""
Best-effort parsers for statement formats beyond the 4 native PDF layouts
handled in parsers.py: CSV/Excel bank exports, Word documents, and photographed
or scanned images (via OCR).

These are deliberately more forgiving than the PDF parsers, because there is
no single fixed layout to target:

- CSV/XLSX: detects which column is the date / description / amount (or
  separate debit+credit columns) by matching common header names, rather
  than assuming a fixed column order. Works well whenever the file has a
  proper header row — which almost every bank's own CSV/Excel export does.
- DOCX: uses the first table found in the document (same column-header
  detection as CSV/XLSX) if there is one; otherwise falls back to scanning
  the document's plain paragraph text line-by-line.
- Images (JPG/PNG): runs OCR (tesseract) to turn the photo/scan into text,
  then applies the same line-by-line fallback as DOCX-without-a-table.

The line-by-line fallback (used for DOCX without a table, and always for
images) is inherently the least reliable path here — there's no structured
column data to rely on, just "does this line look like a date, then some
text, then an amount". It's a reasonable best effort, not a guarantee: a
native PDF or CSV/Excel export from your bank will always parse more
accurately than a photo of a screen or a printed page.
"""
import re
import csv
import datetime
from dateutil import parser as dateutil_parser

from parsers import UnknownStatementType

GENERIC_ACCOUNT_FALLBACK = 'Other/Imported Statement'

DATE_ALIASES = {'date', 'transaction date', 'posting date', 'value date', 'trans date', 'txn date'}
DESC_ALIASES = {'description', 'details', 'narrative', 'reference', 'memo', 'payee', 'merchant',
                'transaction description', 'transaction'}
AMOUNT_ALIASES = {'amount', 'value', 'transaction amount', 'amount (gbp)', 'amount (£)'}
DEBIT_ALIASES = {'debit', 'withdrawal', 'paid out', 'money out', 'debit amount', 'out'}
CREDIT_ALIASES = {'credit', 'deposit', 'paid in', 'money in', 'credit amount', 'in'}
TYPE_ALIASES = {'type', 'transaction type'}
BALANCE_ALIASES = {'balance', 'running balance', 'closing balance'}

ALL_HEADER_ALIASES = (DATE_ALIASES | DESC_ALIASES | AMOUNT_ALIASES | DEBIT_ALIASES |
                       CREDIT_ALIASES | TYPE_ALIASES | BALANCE_ALIASES)


# ---------- shared helpers ----------

def _normalize_cell(c):
    if c is None:
        return ''
    return re.sub(r'\s+', ' ', str(c)).strip().lower()


def _row_is_empty(row):
    return all(_normalize_cell(c) == '' for c in row)


def _find_header_row(rows, max_scan=10):
    """Best-effort: the header row is whichever of the first few rows has the
    most cells matching a known column-name alias (handles CSV/Excel exports
    that have a line or two of preamble before the real header)."""
    best_idx, best_score = 0, -1
    for i, row in enumerate(rows[:max_scan]):
        cells = [_normalize_cell(c) for c in row]
        score = sum(1 for c in cells if c in ALL_HEADER_ALIASES)
        if score > best_score:
            best_score, best_idx = score, i
    return best_idx if best_score >= 2 else 0


def _find_col(headers, aliases):
    for i, h in enumerate(headers):
        if h in aliases:
            return i
    return None


def _to_amount(raw):
    """Parse a money-ish cell/string into a signed float, or None if it isn't
    one. Negative is signalled by a leading '-', wrapping parentheses, or a
    'CR' suffix removed beforehand by the caller — this function itself just
    handles '-' and '(...)' plus stripping '£'/'$'/commas/'CR'/'DR'."""
    if raw is None:
        return None
    if isinstance(raw, (int, float)):
        return float(raw)
    s = str(raw).strip()
    if not s:
        return None
    neg = False
    if s.startswith('(') and s.endswith(')'):
        neg = True
        s = s[1:-1]
    s = s.replace('£', '').replace('$', '').replace(',', '').strip()
    upper = s.upper()
    if upper.endswith('CR'):
        s = s[:-2].strip()
    elif upper.endswith('DR'):
        s = s[:-2].strip()
    if s.startswith('-'):
        neg = True
        s = s[1:].strip()
    if not s:
        return None
    try:
        val = float(s)
    except ValueError:
        return None
    return -val if neg else val


def normalize_date(raw):
    """Turn almost any date representation (a datetime/date object from
    Excel, or a string in whatever format) into the 'DD Mon YY' string the
    rest of the app expects (matches categorization.parse_month's regex)."""
    if raw is None:
        return None
    if isinstance(raw, (datetime.date, datetime.datetime)):
        d = raw
        return f"{d.day:02d} {d.strftime('%b')} {str(d.year)[2:]}"
    s = str(raw).strip()
    if not s:
        return None
    try:
        d = dateutil_parser.parse(s, dayfirst=True, fuzzy=False)
    except (ValueError, OverflowError, TypeError):
        return None
    return f"{d.day:02d} {d.strftime('%b')} {str(d.year)[2:]}"


def _detect_account(blob):
    """Same priority logic as parsers.sniff_header()/detect_and_parse(), just
    run over whatever text we have available (headers, a few data rows,
    filename) instead of a PDF's letterhead."""
    b = blob.upper()
    if 'HALIFAX' in b and ('CURRENT ACCOUNT' in b or 'SORT CODE' in b):
        return 'Halifax Current Account'
    if 'VISA CARD STATEMENT' in b or ('HSBC' in b and 'CREDIT LIMIT' in b):
        return 'HSBC Credit Card'
    if 'BARCLAYCARD' in b:
        return 'Barclaycard'
    if 'HSBC' in b:
        return 'HSBC Current Account'
    return None


def _rows_to_txns(header_row, data_rows, account):
    headers = [_normalize_cell(c) for c in header_row]
    date_i = _find_col(headers, DATE_ALIASES)
    desc_i = _find_col(headers, DESC_ALIASES)
    amount_i = _find_col(headers, AMOUNT_ALIASES)
    debit_i = _find_col(headers, DEBIT_ALIASES)
    credit_i = _find_col(headers, CREDIT_ALIASES)
    type_i = _find_col(headers, TYPE_ALIASES)

    if date_i is None or desc_i is None or (amount_i is None and debit_i is None and credit_i is None):
        return []

    txns = []
    for row in data_rows:
        if date_i >= len(row) or desc_i >= len(row):
            continue
        date_norm = normalize_date(row[date_i])
        if not date_norm:
            continue
        desc = re.sub(r'\s+', ' ', str(row[desc_i] or '')).strip()
        if not desc:
            continue

        amount, direction = None, None
        if debit_i is not None or credit_i is not None:
            debit_val = _to_amount(row[debit_i]) if debit_i is not None and debit_i < len(row) else None
            credit_val = _to_amount(row[credit_i]) if credit_i is not None and credit_i < len(row) else None
            if debit_val:
                amount, direction = abs(debit_val), 'out'
            elif credit_val:
                amount, direction = abs(credit_val), 'in'
            else:
                continue
        elif amount_i is not None and amount_i < len(row):
            val = _to_amount(row[amount_i])
            if val is None:
                continue
            amount, direction = abs(val), ('out' if val < 0 else 'in')
        else:
            continue

        typ = ''
        if type_i is not None and type_i < len(row) and row[type_i] is not None:
            typ = str(row[type_i]).strip()

        txns.append({'account': account, 'date': date_norm, 'desc': desc,
                      'amount': round(amount, 2), 'dir': direction, 'type': typ})
    return txns


# ---------- line-by-line fallback (DOCX without a table, and all images) ----------

_TEXT_LINE_RE = re.compile(
    r'^(?P<date>\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4}|\d{1,2}\s+[A-Za-z]{3,9}\s+\d{2,4})\s+'
    r'(?P<desc>.+?)\s+'
    r'(?P<amt>-?£?\(?[\d,]+\.\d{2}\)?)\s*(?P<crdr>CR|DR)?'
    r'(?:\s+-?£?[\d,]+\.\d{2})?'  # optional trailing running-balance figure, ignored
    r'[,.\s]*$'  # tolerate stray trailing punctuation from imperfect OCR
)


def _parse_text_lines(text, account):
    txns = []
    for raw_line in text.split('\n'):
        line = raw_line.strip()
        if not line:
            continue
        m = _TEXT_LINE_RE.match(line)
        if not m:
            continue
        date_norm = normalize_date(m.group('date'))
        if not date_norm:
            continue
        desc = re.sub(r'\s+', ' ', m.group('desc')).strip()
        if not desc:
            continue
        amt_raw = m.group('amt').strip()
        looks_credit = amt_raw.startswith('-') or amt_raw.startswith('(')
        val = _to_amount(amt_raw)
        if val is None:
            continue
        crdr = m.group('crdr')
        if crdr == 'CR':
            direction = 'in'
        elif crdr == 'DR':
            direction = 'out'
        else:
            direction = 'in' if looks_credit else 'out'
        txns.append({'account': account, 'date': date_norm, 'desc': desc,
                      'amount': round(abs(val), 2), 'dir': direction, 'type': ''})
    return txns


# ---------- entry points, one per file type ----------

def parse_csv(path, filename):
    with open(path, newline='', encoding='utf-8-sig', errors='ignore') as f:
        rows = list(csv.reader(f))
    rows = [r for r in rows if not _row_is_empty(r)]
    if not rows:
        raise UnknownStatementType("This CSV file appears to be empty.")

    header_idx = _find_header_row(rows)
    header_row = rows[header_idx]
    data_rows = rows[header_idx + 1:]

    blob = filename + ' ' + ' '.join(header_row) + ' ' + ' '.join(' '.join(r) for r in data_rows[:5])
    account = _detect_account(blob) or GENERIC_ACCOUNT_FALLBACK

    txns = _rows_to_txns(header_row, data_rows, account)
    if not txns:
        raise UnknownStatementType(
            "Couldn't find recognisable Date/Description/Amount columns in this CSV. Expected header "
            "names like 'Date', 'Description', and 'Amount' (or separate 'Debit'/'Credit' columns)."
        )
    return account, txns


def parse_excel(path, filename):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    ws = wb.worksheets[0]
    rows = [list(row) for row in ws.iter_rows(values_only=True)]
    rows = [r for r in rows if not _row_is_empty(r)]
    if not rows:
        raise UnknownStatementType("This spreadsheet appears to be empty.")

    header_idx = _find_header_row(rows)
    header_row = rows[header_idx]
    data_rows = rows[header_idx + 1:]

    blob_bits = [filename] + [str(c) for c in header_row if c is not None]
    for r in data_rows[:5]:
        blob_bits.extend(str(c) for c in r if c is not None)
    account = _detect_account(' '.join(blob_bits)) or GENERIC_ACCOUNT_FALLBACK

    txns = _rows_to_txns(header_row, data_rows, account)
    if not txns:
        raise UnknownStatementType(
            "Couldn't find recognisable Date/Description/Amount columns in this spreadsheet. Expected "
            "header names like 'Date', 'Description', and 'Amount' (or separate 'Debit'/'Credit' columns)."
        )
    return account, txns


def parse_docx_file(path, filename):
    import docx
    d = docx.Document(path)

    # Account detection looks at the WHOLE document's text (headings/intro
    # paragraphs commonly name the bank/account even when the transactions
    # themselves are in a table below), not just the table's own cells.
    doc_text = '\n'.join(p.text for p in d.paragraphs)

    if d.tables:
        table = d.tables[0]
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rows = [r for r in rows if not _row_is_empty(r)]
        if rows:
            header_idx = _find_header_row(rows)
            header_row = rows[header_idx]
            data_rows = rows[header_idx + 1:]
            blob = filename + ' ' + doc_text[:2000] + ' ' + ' '.join(header_row)
            account = _detect_account(blob) or GENERIC_ACCOUNT_FALLBACK
            txns = _rows_to_txns(header_row, data_rows, account)
            if txns:
                return account, txns

    # No usable table (or the table didn't yield anything) — fall back to
    # scanning the document's own paragraph text line by line. Lower
    # reliability than the table path: there's no column structure to lean on.
    blob = filename + ' ' + doc_text[:2000]
    account = _detect_account(blob) or GENERIC_ACCOUNT_FALLBACK
    txns = _parse_text_lines(doc_text, account)
    if not txns:
        raise UnknownStatementType(
            "Couldn't find any recognisable transaction lines (date + description + amount) in this "
            "Word document, either as a table or as plain text."
        )
    return account, txns


def parse_image_ocr(path, filename):
    from PIL import Image
    import pytesseract
    img = Image.open(path)
    text = pytesseract.image_to_string(img)
    blob = filename + ' ' + text[:2000]
    account = _detect_account(blob) or GENERIC_ACCOUNT_FALLBACK
    txns = _parse_text_lines(text, account)
    if not txns:
        raise UnknownStatementType(
            "Couldn't read any recognisable transaction lines from this image via text recognition (OCR). "
            "Scanned/photographed statements work best as a clear, straight-on, well-lit, high-resolution "
            "photo — a native PDF or CSV/Excel export from your bank will always parse more reliably."
        )
    return account, txns
